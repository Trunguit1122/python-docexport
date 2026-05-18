"""
Sinh mẫu Word theo thể thức văn bản hành chính Nghị định 30/2020/NĐ-CP.

Mặc định tạo đủ bộ 25 mẫu trình bày văn bản hành chính giấy trong
Phụ lục III. Nội dung trong dấu ngoặc vuông là placeholder để người
dùng hoặc agent khác thay thế.

Cài đặt:
    python3 -m pip install python-docx

Chạy nhanh:
    python3 main.py --all
    python3 main.py --type cong_van
    python3 main.py --list
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_NAME = "Times New Roman"
BODY_SIZE = 13
CONTENT_WIDTH_CM = 16.5


@dataclass(frozen=True)
class TemplateSpec:
    key: str
    title: str
    abbreviation: str
    filename: str
    renderer: str
    subtitle: str = "[TRÍCH YẾU NỘI DUNG VĂN BẢN]"
    has_recipient: bool = False


TEMPLATE_SPECS: tuple[TemplateSpec, ...] = (
    TemplateSpec("nghi_quyet", "NGHỊ QUYẾT", "NQ", "01_Nghi_quyet_ca_biet.docx", "resolution"),
    TemplateSpec("quyet_dinh", "QUYẾT ĐỊNH", "QĐ", "02_Quyet_dinh_ca_biet.docx", "decision"),
    TemplateSpec("chi_thi", "CHỈ THỊ", "CT", "03_Chi_thi.docx", "generic"),
    TemplateSpec("quy_che", "QUY CHẾ", "QC", "04_Quy_che.docx", "generic"),
    TemplateSpec("quy_dinh", "QUY ĐỊNH", "QĐ", "05_Quy_dinh.docx", "generic"),
    TemplateSpec("thong_cao", "THÔNG CÁO", "TC", "06_Thong_cao.docx", "generic"),
    TemplateSpec("thong_bao", "THÔNG BÁO", "TB", "07_Thong_bao.docx", "generic"),
    TemplateSpec("huong_dan", "HƯỚNG DẪN", "HD", "08_Huong_dan.docx", "generic"),
    TemplateSpec("chuong_trinh", "CHƯƠNG TRÌNH", "CTr", "09_Chuong_trinh.docx", "generic"),
    TemplateSpec("ke_hoach", "KẾ HOẠCH", "KH", "10_Ke_hoach.docx", "generic"),
    TemplateSpec("phuong_an", "PHƯƠNG ÁN", "PA", "11_Phuong_an.docx", "generic"),
    TemplateSpec("de_an", "ĐỀ ÁN", "ĐA", "12_De_an.docx", "generic"),
    TemplateSpec("du_an", "DỰ ÁN", "DA", "13_Du_an.docx", "generic"),
    TemplateSpec("bao_cao", "BÁO CÁO", "BC", "14_Bao_cao.docx", "generic"),
    TemplateSpec("to_trinh", "TỜ TRÌNH", "TTr", "15_To_trinh.docx", "generic", has_recipient=True),
    TemplateSpec("giay_uy_quyen", "GIẤY ỦY QUYỀN", "GUQ", "16_Giay_uy_quyen.docx", "authorization"),
    TemplateSpec("phieu_gui", "PHIẾU GỬI", "PG", "17_Phieu_gui.docx", "generic", has_recipient=True),
    TemplateSpec("phieu_chuyen", "PHIẾU CHUYỂN", "PC", "18_Phieu_chuyen.docx", "generic", has_recipient=True),
    TemplateSpec("phieu_bao", "PHIẾU BÁO", "PB", "19_Phieu_bao.docx", "generic", has_recipient=True),
    TemplateSpec(
        "cong_van",
        "CÔNG VĂN",
        "CV",
        "20_Cong_van.docx",
        "dispatch",
        subtitle="[TRÍCH YẾU NỘI DUNG CÔNG VĂN]",
        has_recipient=True,
    ),
    TemplateSpec("cong_dien", "CÔNG ĐIỆN", "CĐ", "21_Cong_dien.docx", "telegram", has_recipient=True),
    TemplateSpec("giay_moi", "GIẤY MỜI", "GM", "22_Giay_moi.docx", "invitation", has_recipient=True),
    TemplateSpec("giay_gioi_thieu", "GIẤY GIỚI THIỆU", "GGT", "23_Giay_gioi_thieu.docx", "introduction"),
    TemplateSpec("bien_ban", "BIÊN BẢN", "BB", "24_Bien_ban.docx", "minutes"),
    TemplateSpec("giay_nghi_phep", "GIẤY NGHỈ PHÉP", "GNP", "25_Giay_nghi_phep.docx", "leave"),
)

SPEC_BY_KEY = {spec.key: spec for spec in TEMPLATE_SPECS}


def set_run_font(run, *, size: int = BODY_SIZE, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_run(paragraph, text: str, *, size: int = BODY_SIZE, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return run


def set_style_font(style, *, size: int = BODY_SIZE, bold: bool = False) -> None:
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.bold = bold


def set_paragraph_format(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_cm: float | None = None,
    left_indent_cm: float | None = None,
    space_before_pt: int = 0,
    space_after_pt: int = 6,
    line_spacing: float = 1.15,
) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing = line_spacing
    if first_line_cm is not None:
        fmt.first_line_indent = Cm(first_line_cm)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    size: int = BODY_SIZE,
    bold: bool = False,
    italic: bool = False,
    first_line_cm: float | None = None,
    left_indent_cm: float | None = None,
    space_before_pt: int = 0,
    space_after_pt: int = 6,
):
    paragraph = doc.add_paragraph()
    set_paragraph_format(
        paragraph,
        align=align,
        first_line_cm=first_line_cm,
        left_indent_cm=left_indent_cm,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
    )
    if text:
        add_run(paragraph, text, size=size, bold=bold, italic=italic)
    return paragraph


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def add_short_line(paragraph, *, length: int = 17) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, "-" * length, size=12, bold=True)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=13)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.different_first_page_header_footer = True

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.text = ""
    header_paragraph.paragraph_format.space_after = Pt(0)
    add_page_number(header_paragraph)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=BODY_SIZE)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        set_style_font(doc.styles[style_name], size=BODY_SIZE, bold=True)

    return doc


def prepare_table(table, widths_cm: tuple[float, ...]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def cell_paragraph(cell, text: str = "", *, align=WD_ALIGN_PARAGRAPH.CENTER, size=BODY_SIZE, bold=False, italic=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, align=align, space_after_pt=0)
    if text:
        add_run(paragraph, text, size=size, bold=bold, italic=italic)
    return paragraph


def add_top_block(doc: Document, spec: TemplateSpec) -> None:
    table = doc.add_table(rows=2, cols=2)
    prepare_table(table, (6.6, 9.9))

    left_top = table.cell(0, 0)
    cell_paragraph(left_top, "[CƠ QUAN CHỦ QUẢN]", size=12)
    paragraph = left_top.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0)
    add_run(paragraph, "[CƠ QUAN BAN HÀNH]", size=12, bold=True)
    add_short_line(left_top.add_paragraph(), length=15)

    right_top = table.cell(0, 1)
    cell_paragraph(right_top, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=12, bold=True)
    paragraph = right_top.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0)
    add_run(paragraph, "Độc lập - Tự do - Hạnh phúc", size=13, bold=True)
    add_short_line(right_top.add_paragraph(), length=21)

    left_bottom = table.cell(1, 0)
    cell_paragraph(left_bottom, f"Số: [SỐ]/{spec.abbreviation}-[CQBH]", size=13)
    if spec.renderer == "dispatch":
        paragraph = left_bottom.add_paragraph()
        set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0)
        add_run(paragraph, f"V/v {spec.subtitle}", size=12)

    right_bottom = table.cell(1, 1)
    cell_paragraph(right_bottom, "[Địa danh], ngày [..] tháng [..] năm [....]", size=13, italic=True)
    doc.add_paragraph()


def add_title_block(doc: Document, spec: TemplateSpec, *, include_subtitle: bool = True) -> None:
    add_paragraph(doc, spec.title, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after_pt=0)
    if include_subtitle:
        add_paragraph(doc, spec.subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after_pt=0)
    paragraph = doc.add_paragraph()
    add_short_line(paragraph, length=19)
    paragraph.paragraph_format.space_after = Pt(6)


def add_kinh_gui(doc: Document, *, multiple: bool = False) -> None:
    if multiple:
        add_paragraph(doc, "Kính gửi:", size=13, bold=True, left_indent_cm=2.2, space_after_pt=0)
        for recipient in ("[NƠI NHẬN THỨ NHẤT]", "[NƠI NHẬN THỨ HAI]"):
            add_paragraph(doc, f"- {recipient};", left_indent_cm=3.1, space_after_pt=0)
        return
    add_paragraph(
        doc,
        "Kính gửi: [TÊN CƠ QUAN/TỔ CHỨC/CÁ NHÂN NHẬN]",
        size=13,
        bold=True,
        left_indent_cm=2.2,
        space_after_pt=6,
    )


def add_body_lines(doc: Document, lines: tuple[str, ...]) -> None:
    for line in lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=1.0, space_after_pt=6)


def add_generic_body(doc: Document, spec: TemplateSpec) -> None:
    if spec.has_recipient:
        add_kinh_gui(doc)
    add_body_lines(
        doc,
        (
            "[Nhập phần mở đầu/căn cứ/mục đích ban hành văn bản.]",
            "[Nhập nội dung chính. Nếu văn bản có mục, khoản, điều thì giữ cách đánh số thống nhất.]",
            "[Nhập phần kết luận, yêu cầu thực hiện hoặc đề nghị phối hợp nếu có.]",
        ),
    )


def add_legal_basis(doc: Document) -> None:
    for line in (
        "Căn cứ [tên văn bản quy định thẩm quyền/chức năng/nhiệm vụ];",
        "Căn cứ [văn bản/kết luận/đề nghị liên quan];",
        "Theo đề nghị của [đơn vị/cá nhân tham mưu],",
    ):
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=1.0, italic=True, space_after_pt=6)


def add_decision_body(doc: Document) -> None:
    add_legal_basis(doc)
    add_paragraph(doc, "QUYẾT ĐỊNH:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after_pt=6)
    add_body_lines(
        doc,
        (
            "Điều 1. [Nội dung quyết định.]",
            "Điều 2. [Trách nhiệm thi hành của các cơ quan, đơn vị, cá nhân liên quan.]",
            "Điều 3. [Hiệu lực thi hành và tổ chức thực hiện.]",
        ),
    )


def add_resolution_body(doc: Document) -> None:
    add_legal_basis(doc)
    add_paragraph(doc, "QUYẾT NGHỊ:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after_pt=6)
    add_body_lines(
        doc,
        (
            "Điều 1. [Nội dung nghị quyết.]",
            "Điều 2. [Tổ chức thực hiện.]",
            "Điều 3. [Hiệu lực và trách nhiệm thi hành.]",
        ),
    )


def add_dispatch_body(doc: Document) -> None:
    add_kinh_gui(doc, multiple=True)
    add_body_lines(
        doc,
        (
            "[Nêu lý do, bối cảnh hoặc căn cứ gửi công văn.]",
            "[Nội dung trao đổi/đề nghị/hướng dẫn/trả lời.]",
            "[Đề nghị quý cơ quan/đơn vị phối hợp thực hiện.]",
        ),
    )


def add_telegram_body(doc: Document) -> None:
    add_kinh_gui(doc, multiple=True)
    add_body_lines(
        doc,
        (
            "[Nhập nội dung chỉ đạo/thông tin khẩn của công điện.]",
            "[Yêu cầu về thời hạn, đơn vị thực hiện và chế độ báo cáo.]",
        ),
    )


def add_invitation_body(doc: Document) -> None:
    add_kinh_gui(doc)
    for label in (
        "Trân trọng kính mời: [TÊN CƠ QUAN/TỔ CHỨC/CÁ NHÂN]",
        "Dự họp/dự [TÊN SỰ KIỆN]: [NỘI DUNG]",
        "Thời gian: [GIỜ], ngày [..] tháng [..] năm [....]",
        "Địa điểm: [ĐỊA ĐIỂM]",
        "Thành phần: [THÀNH PHẦN DỰ]",
        "Đề nghị: [YÊU CẦU CHUẨN BỊ NẾU CÓ]",
    ):
        add_paragraph(doc, label, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=1.0, space_after_pt=6)


def add_introduction_body(doc: Document) -> None:
    add_body_lines(
        doc,
        (
            "Giới thiệu ông/bà: [HỌ VÀ TÊN]; chức vụ: [CHỨC VỤ]; đơn vị: [ĐƠN VỊ].",
            "Được cử đến: [CƠ QUAN/TỔ CHỨC LIÊN HỆ].",
            "Về việc: [NỘI DUNG LIÊN HỆ/CÔNG TÁC].",
            "Giấy này có giá trị đến hết ngày [..] tháng [..] năm [....].",
        ),
    )


def add_authorization_body(doc: Document) -> None:
    add_body_lines(
        doc,
        (
            "Bên ủy quyền: [HỌ TÊN/CHỨC VỤ/ĐƠN VỊ].",
            "Bên được ủy quyền: [HỌ TÊN/CHỨC VỤ/ĐƠN VỊ].",
            "Nội dung ủy quyền: [PHẠM VI CÔNG VIỆC ĐƯỢC ỦY QUYỀN].",
            "Thời hạn ủy quyền: từ ngày [..] tháng [..] năm [....] đến ngày [..] tháng [..] năm [....].",
            "Các bên chịu trách nhiệm thực hiện đúng nội dung ủy quyền nêu trên.",
        ),
    )


def add_leave_body(doc: Document) -> None:
    add_body_lines(
        doc,
        (
            "Cấp cho ông/bà: [HỌ VÀ TÊN]; chức vụ: [CHỨC VỤ]; đơn vị: [ĐƠN VỊ].",
            "Được nghỉ phép từ ngày [..] tháng [..] năm [....] đến hết ngày [..] tháng [..] năm [....].",
            "Nơi nghỉ phép: [ĐỊA CHỈ NƠI NGHỈ PHÉP].",
            "Lý do nghỉ phép: [LÝ DO].",
            "Ông/bà [HỌ VÀ TÊN] có trách nhiệm bàn giao công việc và trở lại làm việc đúng thời hạn.",
        ),
    )


def add_minutes_body(doc: Document) -> None:
    for line in (
        "Thời gian bắt đầu: [GIỜ] giờ [PHÚT], ngày [..] tháng [..] năm [....].",
        "Địa điểm: [ĐỊA ĐIỂM LẬP BIÊN BẢN].",
        "Thành phần tham dự: [DANH SÁCH THÀNH PHẦN].",
        "Nội dung: [NỘI DUNG SỰ VIỆC/CUỘC HỌP/TRAO ĐỔI ĐƯỢC GHI NHẬN].",
        "Biên bản kết thúc vào lúc [GIỜ] giờ [PHÚT] cùng ngày, được lập thành [SỐ] bản có giá trị như nhau.",
    ):
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=1.0, space_after_pt=6)


def add_signature_block(doc: Document, *, minutes: bool = False) -> None:
    table = doc.add_table(rows=1, cols=2)
    prepare_table(table, (7.6, 8.9))

    left = table.cell(0, 0)
    if minutes:
        cell_paragraph(left, "ĐẠI DIỆN CÁC BÊN", bold=True)
        for _ in range(5):
            left.add_paragraph()
        paragraph = left.add_paragraph()
        set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0)
        add_run(paragraph, "[HỌ VÀ TÊN]", bold=True)
    else:
        cell_paragraph(left, "Nơi nhận:", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True, italic=True)
        for idx, item in enumerate(("[Như trên]", "[Cơ quan/đơn vị liên quan]", "Lưu: VT, [BỘ PHẬN SOẠN]")):
            suffix = ";" if idx < 2 else "."
            paragraph = left.add_paragraph()
            set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=0)
            add_run(paragraph, f"- {item}{suffix}", size=11)

    right = table.cell(0, 1)
    heading = "NGƯỜI LẬP BIÊN BẢN" if minutes else "[THẨM QUYỀN KÝ]"
    cell_paragraph(right, heading, bold=True)
    paragraph = right.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0)
    add_run(paragraph, "[CHỨC VỤ NGƯỜI KÝ]", bold=True)

    for _ in range(5):
        right.add_paragraph()

    paragraph = right.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0)
    add_run(paragraph, "[HỌ VÀ TÊN]", bold=True)


def add_contact_line(doc: Document) -> None:
    add_paragraph(
        doc,
        "[Thông tin liên hệ nếu cần: địa chỉ, điện thoại, email]",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        space_before_pt=6,
        space_after_pt=0,
    )


def render_named(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec)
    add_generic_body(doc, spec)
    add_signature_block(doc)
    add_contact_line(doc)


def render_decision(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec)
    add_decision_body(doc)
    add_signature_block(doc)


def render_resolution(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec)
    add_resolution_body(doc)
    add_signature_block(doc)


def render_dispatch(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_dispatch_body(doc)
    add_signature_block(doc)
    add_contact_line(doc)


def render_telegram(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec, include_subtitle=False)
    add_telegram_body(doc)
    add_signature_block(doc)


def render_invitation(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec, include_subtitle=False)
    add_invitation_body(doc)
    add_signature_block(doc)


def render_introduction(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec, include_subtitle=False)
    add_introduction_body(doc)
    add_signature_block(doc)


def render_authorization(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec, include_subtitle=False)
    add_authorization_body(doc)
    add_signature_block(doc)


def render_leave(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec, include_subtitle=False)
    add_leave_body(doc)
    add_signature_block(doc)


def render_minutes(doc: Document, spec: TemplateSpec) -> None:
    add_top_block(doc, spec)
    add_title_block(doc, spec)
    add_minutes_body(doc)
    add_signature_block(doc, minutes=True)


RENDERERS: dict[str, Callable[[Document, TemplateSpec], None]] = {
    "generic": render_named,
    "decision": render_decision,
    "resolution": render_resolution,
    "dispatch": render_dispatch,
    "telegram": render_telegram,
    "invitation": render_invitation,
    "introduction": render_introduction,
    "authorization": render_authorization,
    "leave": render_leave,
    "minutes": render_minutes,
}


def create_template(spec_key: str, path: str | Path) -> Path:
    if spec_key not in SPEC_BY_KEY:
        available = ", ".join(sorted(SPEC_BY_KEY))
        raise ValueError(f"Không hỗ trợ loại văn bản '{spec_key}'. Các loại hợp lệ: {available}")

    spec = SPEC_BY_KEY[spec_key]
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = setup_document()
    RENDERERS[spec.renderer](doc, spec)
    doc.save(path)
    return path


def create_all_templates(output_dir: str | Path) -> list[Path]:
    output_dir = Path(output_dir)
    return [create_template(spec.key, output_dir / spec.filename) for spec in TEMPLATE_SPECS]


def list_template_types() -> str:
    width = max(len(spec.key) for spec in TEMPLATE_SPECS)
    return "\n".join(f"{spec.key:<{width}}  ->  {spec.filename}" for spec in TEMPLATE_SPECS)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Tạo mẫu Word theo Nghị định 30/2020/NĐ-CP.")
    parser.add_argument("--out-dir", default="output_nd30", help="Thư mục xuất file. Mặc định: output_nd30")
    parser.add_argument("--type", choices=sorted(SPEC_BY_KEY), help="Loại văn bản cần tạo, ví dụ: cong_van")
    parser.add_argument("--all", action="store_true", help="Tạo tất cả 25 mẫu.")
    parser.add_argument("--list", action="store_true", help="Liệt kê các loại văn bản hỗ trợ.")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    if args.list:
        print(list_template_types())
        return 0

    output_dir = Path(args.out_dir)
    if args.type:
        spec = SPEC_BY_KEY[args.type]
        path = create_template(args.type, output_dir / spec.filename)
        print(f"Đã tạo: {path.resolve()}")
        return 0

    paths = create_all_templates(output_dir)
    print(f"Đã tạo {len(paths)} file trong: {output_dir.resolve()}")
    for path in paths:
        print(f"- {path.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
